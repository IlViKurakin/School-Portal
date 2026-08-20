
from __future__ import annotations

import html
import io
import json
import re
import sqlite3
import uuid

from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document

from flask import (
    abort,
    flash,
    g,
    redirect,
    request,
    send_file,
    url_for,
)


# ============================================================
# Справочники
# ============================================================

RU_MONTHS = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}


def apply_post_enrollment_v8(
    app,
    namespace: dict[str, Any],
) -> None:

    base_dir: Path = namespace["BASE_DIR"]
    database_path: Path = namespace["DATABASE_PATH"]

    get_db = namespace["get_db"]
    get_student_or_404 = namespace[
        "get_student_or_404"
    ]

    login_required = namespace["login_required"]
    roles_required = namespace["roles_required"]

    audit = namespace["audit"]
    render_page = namespace["render_page"]
    full_name = namespace["full_name"]

    template_dir = base_dir / "templates"

    order_template = (
        template_dir
        / "order_extract_template.docx"
    )

    certificate_template = (
        template_dir
        / "certificate_template.docx"
    )

    generated_dir = (
        base_dir
        / "generated_documents"
    )

    generated_dir.mkdir(
        exist_ok=True
    )

    # ========================================================
    # База данных
    # ========================================================

    def default_academic_year() -> str:
        today = date.today()

        # Июнь и далее уже считаем
        # следующим учебным годом.
        if today.month >= 6:
            return (
                f"{today.year}-"
                f"{today.year + 1}"
            )

        return (
            f"{today.year - 1}-"
            f"{today.year}"
        )

    def migrate_database() -> None:

        connection = sqlite3.connect(
            database_path
        )

        try:
            connection.execute(
                "PRAGMA foreign_keys = ON"
            )

            connection.executescript(
                """
                CREATE TABLE IF NOT EXISTS
                post_enrollment_documents (
                    id INTEGER
                        PRIMARY KEY AUTOINCREMENT,

                    student_id INTEGER NOT NULL,

                    document_type TEXT NOT NULL
                        CHECK (
                            document_type IN (
                                'order_extract',
                                'certificate'
                            )
                        ),

                    status TEXT NOT NULL
                        DEFAULT 'draft'
                        CHECK (
                            status IN (
                                'draft',
                                'issued'
                            )
                        ),

                    certificate_number INTEGER,

                    order_number TEXT,
                    order_date TEXT,

                    issue_date TEXT,
                    academic_year TEXT,

                    name_override TEXT,

                    snapshot_json TEXT,
                    stored_name TEXT,

                    created_by INTEGER NOT NULL,
                    created_at TEXT NOT NULL,

                    updated_by INTEGER,
                    updated_at TEXT NOT NULL,

                    finalized_by INTEGER,
                    finalized_at TEXT,

                    FOREIGN KEY (student_id)
                        REFERENCES students(id),

                    FOREIGN KEY (created_by)
                        REFERENCES users(id),

                    FOREIGN KEY (updated_by)
                        REFERENCES users(id),

                    FOREIGN KEY (finalized_by)
                        REFERENCES users(id),

                    UNIQUE (
                        student_id,
                        document_type
                    ),

                    UNIQUE (
                        document_type,
                        certificate_number
                    )
                );

                CREATE TABLE IF NOT EXISTS
                document_counters (
                    counter_key TEXT
                        PRIMARY KEY,

                    last_value INTEGER
                        NOT NULL DEFAULT 0
                );

                CREATE TABLE IF NOT EXISTS
                portal_settings (
                    setting_key TEXT
                        PRIMARY KEY,

                    setting_value TEXT
                        NOT NULL
                );

                CREATE INDEX IF NOT EXISTS
                idx_post_docs_student
                ON post_enrollment_documents(
                    student_id
                );

                CREATE INDEX IF NOT EXISTS
                idx_post_docs_type
                ON post_enrollment_documents(
                    document_type
                );

                CREATE UNIQUE INDEX IF NOT EXISTS
                idx_certificate_number
                ON post_enrollment_documents(
                    certificate_number
                )
                WHERE
                    document_type = 'certificate'
                    AND certificate_number
                        IS NOT NULL;
                """
            )

            connection.execute(
                """
                INSERT OR IGNORE INTO
                    document_counters (
                        counter_key,
                        last_value
                    )
                VALUES (
                    'certificate',
                    0
                )
                """
            )

            connection.execute(
                """
                INSERT OR IGNORE INTO
                    portal_settings (
                        setting_key,
                        setting_value
                    )
                VALUES (
                    'academic_year',
                    ?
                )
                """,
                (
                    default_academic_year(),
                ),
            )

            # Если документы уже когда-нибудь
            # появятся в таблице, счетчик
            # никогда не должен оказаться
            # ниже существующего максимума.
            maximum = connection.execute(
                """
                SELECT COALESCE(
                    MAX(certificate_number),
                    0
                )
                FROM post_enrollment_documents
                WHERE document_type =
                    'certificate'
                """
            ).fetchone()[0]

            connection.execute(
                """
                UPDATE document_counters
                SET last_value =
                    MAX(last_value, ?)
                WHERE counter_key =
                    'certificate'
                """,
                (maximum,),
            )

            connection.commit()

        finally:
            connection.close()

    migrate_database()

    # ========================================================
    # Общие функции
    # ========================================================

    def now_iso() -> str:
        return datetime.now().isoformat(
            timespec="seconds"
        )

    def academic_year() -> str:
        row = get_db().execute(
            """
            SELECT setting_value
            FROM portal_settings
            WHERE setting_key =
                'academic_year'
            """
        ).fetchone()

        if row:
            return row[0]

        return default_academic_year()

    def date_from_iso(
        value: str,
    ) -> date:
        return date.fromisoformat(
            value
        )

    def short_date(
        value: str | date,
    ) -> str:

        parsed = (
            value
            if isinstance(value, date)
            else date_from_iso(value)
        )

        return parsed.strftime(
            "%d.%m.%Y"
        )

    def russian_date(
        value: str | date,
        ending: str = "г.",
    ) -> str:

        parsed = (
            value
            if isinstance(value, date)
            else date_from_iso(value)
        )

        return (
            f"«{parsed.day:02d}» "
            f"{RU_MONTHS[parsed.month]} "
            f"{parsed.year}{ending}"
        )

    def student_full_name(
        student,
    ) -> str:

        return full_name(
            student["last_name"],
            student["first_name"],
            student["middle_name"],
        )

    def require_enrolled(
        student,
    ) -> None:

        if student["status"] != "enrolled":
            abort(
                400,
                "Документы после зачисления "
                "можно формировать только "
                "для ученика со статусом "
                "«Зачислен»."
            )

    def get_post_document(
        student_id: int,
        document_type: str,
    ):

        return get_db().execute(
            """
            SELECT *
            FROM post_enrollment_documents
            WHERE student_id = ?
              AND document_type = ?
            """,
            (
                student_id,
                document_type,
            ),
        ).fetchone()

    # ========================================================
    # Дательный падеж ФИО для справки
    #
    # Это вспомогательная автоматизация.
    # На экране предварительного просмотра
    # сотрудник может исправить форму ФИО.
    # ========================================================

    def dative_surname(
        value: str,
        gender: str,
    ) -> str:

        value = value.strip()

        if not value:
            return ""

        lower = value.lower()

        female = (
            str(gender).lower()
            .startswith("жен")
        )

        if female:

            if lower.endswith("ская"):
                return (
                    value[:-4]
                    + "ской"
                )

            if lower.endswith("цкая"):
                return (
                    value[:-4]
                    + "цкой"
                )

            if lower.endswith(
                (
                    "ова",
                    "ева",
                    "ина",
                    "ына",
                )
            ):
                return (
                    value[:-1]
                    + "ой"
                )

            # Неподходящие или
            # несклоняемые фамилии
            # оставляем без изменения.
            return value

        if lower.endswith("ский"):
            return (
                value[:-4]
                + "скому"
            )

        if lower.endswith("цкий"):
            return (
                value[:-4]
                + "цкому"
            )

        if lower.endswith(
            (
                "ов",
                "ев",
                "ин",
                "ын",
            )
        ):
            return value + "у"

        return value

    def dative_first_name(
        value: str,
        gender: str,
    ) -> str:

        value = value.strip()

        if not value:
            return ""

        lower = value.lower()

        female = (
            str(gender).lower()
            .startswith("жен")
        )

        if female:

            if lower.endswith("ия"):
                return (
                    value[:-2]
                    + "ии"
                )

            if lower.endswith(
                ("а", "я")
            ):
                return (
                    value[:-1]
                    + "е"
                )

            return value

        if lower.endswith("а"):
            return (
                value[:-1]
                + "е"
            )

        if lower.endswith(
            ("й", "ь")
        ):
            return (
                value[:-1]
                + "ю"
            )

        if re.search(
            r"[бвгджзклмнпрстфхцчшщ]$",
            lower,
        ):
            return value + "у"

        return value

    def dative_middle_name(
        value: str,
        gender: str,
    ) -> str:

        value = value.strip()

        if not value:
            return ""

        lower = value.lower()

        female = (
            str(gender).lower()
            .startswith("жен")
        )

        if female and lower.endswith("на"):
            return (
                value[:-1]
                + "е"
            )

        if (
            not female
            and lower.endswith("ич")
        ):
            return value + "у"

        return value

    def automatic_certificate_name(
        student,
    ) -> str:

        gender = (
            student["gender"]
            or ""
        )

        return " ".join(
            part
            for part in (
                dative_surname(
                    student["last_name"],
                    gender,
                ),

                dative_first_name(
                    student["first_name"],
                    gender,
                ),

                dative_middle_name(
                    student[
                        "middle_name"
                    ]
                    or "",
                    gender,
                ),
            )
            if part
        )

    # ========================================================
    # Работа с Word
    # ========================================================

    def iter_paragraphs(
        document,
    ):
        for paragraph in document.paragraphs:
            yield paragraph

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in (
                        cell.paragraphs
                    ):
                        yield paragraph

    def replace_span(
        paragraph,
        start: int,
        end: int,
        replacement: str,
    ) -> None:

        runs = paragraph.runs

        if not runs:
            paragraph.add_run(
                replacement
            )
            return

        positions = []

        cursor = 0

        for index, run in enumerate(runs):
            run_start = cursor
            run_end = (
                cursor
                + len(run.text)
            )

            positions.append(
                (
                    index,
                    run_start,
                    run_end,
                )
            )

            cursor = run_end

        first = None
        last = None

        for item in positions:
            index, run_start, run_end = item

            if (
                first is None
                and start < run_end
            ):
                first = item

            if end <= run_end:
                last = item
                break

        if first is None:
            return

        if last is None:
            last = positions[-1]

        first_index = first[0]
        first_start = first[1]

        last_index = last[0]
        last_start = last[1]

        first_offset = (
            start
            - first_start
        )

        last_offset = (
            end
            - last_start
        )

        if first_index == last_index:

            text = runs[
                first_index
            ].text

            runs[
                first_index
            ].text = (
                text[:first_offset]
                + replacement
                + text[last_offset:]
            )

            return

        first_text = runs[
            first_index
        ].text

        last_text = runs[
            last_index
        ].text

        runs[
            first_index
        ].text = (
            first_text[:first_offset]
            + replacement
        )

        for index in range(
            first_index + 1,
            last_index,
        ):
            runs[index].text = ""

        runs[
            last_index
        ].text = (
            last_text[last_offset:]
        )

    def replace_literal(
        paragraph,
        old: str,
        new: str,
    ) -> bool:

        text = paragraph.text

        position = text.find(old)

        if position < 0:
            return False

        replace_span(
            paragraph,
            position,
            position + len(old),
            new,
        )

        return True

    def replace_literal_all(
        paragraph,
        old: str,
        new: str,
    ) -> None:

        while replace_literal(
            paragraph,
            old,
            new,
        ):
            pass

    def replace_regex(
        paragraph,
        pattern: str,
        replacement: str,
        flags: int = 0,
    ) -> bool:

        match = re.search(
            pattern,
            paragraph.text,
            flags=flags,
        )

        if not match:
            return False

        replace_span(
            paragraph,
            match.start(),
            match.end(),
            replacement,
        )

        return True

    def delete_paragraph(
        paragraph,
    ) -> None:

        element = paragraph._element
        parent = element.getparent()

        if parent is not None:
            parent.remove(element)

    def document_to_bytes(
        document,
    ) -> bytes:

        output = io.BytesIO()

        document.save(output)

        return output.getvalue()

    # ========================================================
    # Выписка из приказа
    # ========================================================

    def generate_order_extract(
        snapshot: dict[str, Any],
    ) -> bytes:

        if not order_template.exists():
            raise RuntimeError(
                "Не найден шаблон "
                "order_extract_template.docx."
            )

        document = Document(
            str(order_template)
        )

        order_date = date_from_iso(
            snapshot["order_date"]
        )

        issue_date = date_from_iso(
            snapshot["issue_date"]
        )

        order_heading = (
            "от "
            + russian_date(
                order_date,
                "г.",
            )
            + " № "
            + snapshot["order_number"]
        )

        issue_text = russian_date(
            issue_date,
            "г.",
        )

        paragraphs = list(
            iter_paragraphs(
                document
            )
        )

        for paragraph in paragraphs:

            text = paragraph.text

            # Дата и номер исходного приказа.
            if (
                "от «__»" in text
                and "№" in text
            ):
                replace_regex(
                    paragraph,
                    (
                        r"от\s+«__»"
                        r".*?"
                        r"№\s*"
                        r"___-Э-26/27"
                    ),
                    order_heading,
                )

            # ФИО ученика.
            if (
                "Обучающийся:"
                in paragraph.text
                and "_"
                in paragraph.text
            ):
                replace_regex(
                    paragraph,
                    (
                        r"Обучающийся:"
                        r"\s*_+"
                    ),
                    (
                        "Обучающийся: "
                        + snapshot[
                            "student_name"
                        ]
                    ),
                )

            # Вторая строка
            # старого поля ФИО.
            if (
                "(ФИО полностью)"
                in paragraph.text
            ):
                replace_regex(
                    paragraph,
                    (
                        r"_+\s*"
                        r"\(ФИО полностью\)"
                    ),
                    "",
                )

                if not paragraph.text.strip():
                    delete_paragraph(
                        paragraph
                    )
                    continue

            # Класс.
            if re.search(
                r"в\s+_+\s+класс",
                paragraph.text,
                flags=re.IGNORECASE,
            ):
                replace_regex(
                    paragraph,
                    r"в\s+_+\s+класс",
                    (
                        "в "
                        + str(
                            snapshot[
                                "class_number"
                            ]
                        )
                        + " класс"
                    ),
                    flags=re.IGNORECASE,
                )

            # Дата подтверждения
            # выписки внизу.
            if (
                "«____»"
                in paragraph.text
                or re.search(
                    r"«_+»\s+_+\s+202_+",
                    paragraph.text,
                )
            ):
                replace_regex(
                    paragraph,
                    (
                        r"«_+»"
                        r"\s+_+"
                        r"\s+202_+г\."
                    ),
                    issue_text,
                )

        return document_to_bytes(
            document
        )

    # ========================================================
    # Справка
    # ========================================================


    # V8.1 CERTIFICATE WITH ORDER REQUISITES
    def generate_certificate(
        snapshot: dict[str, Any],
    ) -> bytes:

        if not certificate_template.exists():
            raise RuntimeError(
                "Не найден шаблон "
                "certificate_template.docx."
            )

        document = Document(
            str(certificate_template)
        )

        issue_date = date_from_iso(
            snapshot["issue_date"]
        )

        certificate_number = str(
            snapshot[
                "certificate_number"
            ]
        )

        birth_date = short_date(
            snapshot["birth_date"]
        )

        order_date_text = short_date(
            snapshot["order_date"]
        )

        # Дата внизу справки —
        # именно дата ее окончательного
        # формирования.
        issue_text = russian_date(
            issue_date,
            " г.",
        )

        for paragraph in list(
            iter_paragraphs(
                document
            )
        ):

            text = paragraph.text

            # ----------------------------------
            # Номер справки
            # Справка № __
            # ----------------------------------

            if re.search(
                r"Справка\s*№\s*_+",
                text,
                flags=re.IGNORECASE,
            ):
                replace_regex(
                    paragraph,
                    r"Справка\s*№\s*_+",
                    (
                        "Справка № "
                        + certificate_number
                    ),
                    flags=re.IGNORECASE,
                )

            # ----------------------------------
            # ФИО + дата рождения
            #
            # Дана ___ФИО полностью___
            # ___ года рождения...
            # ----------------------------------

            if (
                "ФИО полностью"
                in paragraph.text
            ):
                replace_regex(
                    paragraph,
                    (
                        r"Дана\s+"
                        r".*?"
                        r"года рождения"
                    ),
                    (
                        "Дана "
                        + snapshot[
                            "certificate_name"
                        ]
                        + ", "
                        + birth_date
                        + " года рождения"
                    ),
                    flags=re.IGNORECASE,
                )

            # ----------------------------------
            # Класс
            # ----------------------------------

            if re.search(
                (
                    r"зачислен\(а\)"
                    r"\s+в\s+_+"
                    r"\s+класс"
                ),
                paragraph.text,
                flags=re.IGNORECASE,
            ):
                replace_regex(
                    paragraph,
                    (
                        r"зачислен\(а\)"
                        r"\s+в\s+_+"
                        r"\s+класс"
                    ),
                    (
                        "зачислен(а) в "
                        + str(
                            snapshot[
                                "class_number"
                            ]
                        )
                        + " класс"
                    ),
                    flags=re.IGNORECASE,
                )

            # ----------------------------------
            # Учебный год
            #
            # Поддерживаем и
            # 2026-2027,
            # и 2026 - 2027
            # ----------------------------------

            replace_regex(
                paragraph,
                (
                    r"\b20\d{2}"
                    r"\s*-\s*"
                    r"20\d{2}"
                    r"\s+учебном году"
                ),
                (
                    snapshot[
                        "academic_year"
                    ]
                    + " учебном году"
                ),
            )

            # ----------------------------------
            # Номер и дата приказа
            #
            # Приказ о зачислении
            # № ___ от __________.
            # ----------------------------------

            if (
                "Приказ о зачислении"
                in paragraph.text
            ):
                replace_regex(
                    paragraph,
                    (
                        r"Приказ\s+о\s+"
                        r"зачислении"
                        r"\s*№\s*_+"
                        r"\s*от\s*_+"
                        r"\.?"
                    ),
                    (
                        "Приказ о зачислении "
                        "№ "
                        + snapshot[
                            "order_number"
                        ]
                        + " от "
                        + order_date_text
                    ),
                    flags=re.IGNORECASE,
                )

            # ----------------------------------
            # Нижняя дата справки
            #
            # «__» ______ 20___ г.
            # ----------------------------------

            if re.search(
                (
                    r"«_+»"
                    r"\s+_+"
                    r"\s+20_+"
                    r"\s*г\."
                ),
                paragraph.text,
            ):
                replace_regex(
                    paragraph,
                    (
                        r"«_+»"
                        r"\s+_+"
                        r"\s+20_+"
                        r"\s*г\."
                    ),
                    issue_text,
                )

        return document_to_bytes(
            document
        )


    # ========================================================
    # Блок в карточке ученика
    # ========================================================

    def build_student_card_block(
        student,
    ) -> str:

        order_doc = get_post_document(
            student["id"],
            "order_extract",
        )

        certificate_doc = (
            get_post_document(
                student["id"],
                "certificate",
            )
        )

        enrolled = (
            student["status"]
            == "enrolled"
        )

        if not enrolled:

            state = (
                '<span class="status gray">'
                'Доступно после зачисления'
                '</span>'
            )

        else:
            state = (
                '<span class="status green">'
                'Доступно'
                '</span>'
            )

        order_state = (
            "Не сформирована"
        )

        if order_doc:

            if (
                order_doc["status"]
                == "issued"
            ):
                order_state = (
                    "Сформирована "
                    + short_date(
                        order_doc[
                            "issue_date"
                        ]
                    )
                )

            else:
                order_state = "Черновик"

        certificate_state = (
            "Не сформирована"
        )

        if certificate_doc:

            if (
                certificate_doc["status"]
                == "issued"
            ):
                certificate_state = (
                    "Справка № "
                    + str(
                        certificate_doc[
                            "certificate_number"
                        ]
                    )
                )

            else:
                certificate_state = (
                    "Черновик"
                )

        link = url_for(
            "post_docs_v8",
            student_id=student["id"],
        )

        return f"""
        <div
            class="card"
            style="
                margin-top:24px;
                border-top:5px solid #ffd500;
            "
        >
            <div
                style="
                    display:flex;
                    justify-content:
                        space-between;
                    gap:15px;
                    align-items:center;
                    flex-wrap:wrap;
                "
            >
                <div>
                    <h2
                        style="
                            margin:
                            0 0 8px;
                        "
                    >
                        Документы после
                        зачисления
                    </h2>

                    <div class="muted">
                        Выписка из приказа:
                        <strong>
                            {
                                html.escape(
                                    order_state
                                )
                            }
                        </strong>
                        ·
                        Справка:
                        <strong>
                            {
                                html.escape(
                                    certificate_state
                                )
                            }
                        </strong>
                    </div>
                </div>

                <div>
                    {state}

                    <a
                        class="btn btn-primary"
                        style="
                            margin-left:10px;
                        "
                        href="{link}"
                    >
                        Открыть
                    </a>
                </div>
            </div>
        </div>
        """

    original_student_detail = (
        app.view_functions[
            "student_detail"
        ]
    )

    def student_detail_v8(
        student_id: int,
    ):
        result = original_student_detail(
            student_id
        )

        # Redirect / Response не меняем.
        if not isinstance(
            result,
            str,
        ):
            return result

        student = get_student_or_404(
            student_id
        )

        block = (
            build_student_card_block(
                student
            )
        )

        marker = "<h2>Документы</h2>"

        if marker in result:
            return result.replace(
                marker,
                block + marker,
                1,
            )

        # Запасной вариант.
        return result.replace(
            "</section>",
            block + "</section>",
            1,
        )

    app.view_functions[
        "student_detail"
    ] = student_detail_v8

    # ========================================================
    # Главная страница документов после зачисления
    # ========================================================

    def post_docs_page(
        student_id: int,
    ):

        student = get_student_or_404(
            student_id
        )

        order_doc = get_post_document(
            student_id,
            "order_extract",
        )

        certificate_doc = (
            get_post_document(
                student_id,
                "certificate",
            )
        )

        current_year = academic_year()

        body = """
        <div
            style="
                display:flex;
                justify-content:
                    space-between;
                gap:20px;
                align-items:center;
                flex-wrap:wrap;
            "
        >
            <div>
                <h1>
                    Документы после зачисления
                </h1>

                <p class="muted">
                    {{ student.last_name }}
                    {{ student.first_name }}
                    {{ student.middle_name or '' }}
                    ·
                    {{ student.class_number }}
                    класс
                </p>
            </div>

            <a
                class="btn btn-secondary"
                href="{{ url_for(
                    'student_detail',
                    student_id=student.id
                ) }}"
            >
                Назад к карточке
            </a>
        </div>

        {% if student.status != 'enrolled' %}
            <div class="alert info">
                Формирование документов
                станет доступно после
                присвоения карточке статуса
                <strong>«Зачислен»</strong>.
            </div>
        {% endif %}

        <div class="grid metrics">

            <div class="card">
                <h2 style="margin-top:0">
                    Выписка из приказа
                </h2>

                <p>
                    Собственного номера
                    выписка не имеет.
                </p>

                {% if order_doc
                    and
                    order_doc.status
                    == 'issued' %}

                    <p>
                        <span
                            class="status green"
                        >
                            Сформирована
                        </span>
                    </p>

                    <p>
                        Приказ №
                        <strong>
                            {{
                                order_doc.order_number
                            }}
                        </strong>
                        от
                        {{
                            order_doc.order_date
                        }}
                    </p>

                    <p>
                        Дата выписки:
                        <strong>
                            {{
                                order_doc.issue_date
                            }}
                        </strong>
                    </p>

                    <a
                        class="btn btn-primary"
                        href="{{ url_for(
                            'download_post_doc_v8',
                            document_id=
                                order_doc.id
                        ) }}"
                    >
                        Скачать DOCX
                    </a>

                {% elif
                    student.status
                    == 'enrolled' %}

                    {% if order_doc %}
                        <p>
                            <span
                                class="
                                    status yellow
                                "
                            >
                                Черновик
                            </span>
                        </p>
                    {% endif %}

                    <a
                        class="btn btn-primary"
                        href="{{ url_for(
                            'order_extract_v8',
                            student_id=
                                student.id
                        ) }}"
                    >
                        {% if order_doc %}
                            Продолжить
                        {% else %}
                            Сформировать
                        {% endif %}
                    </a>

                {% endif %}
            </div>

            <div class="card">
                <h2 style="margin-top:0">
                    Справка по месту
                    требования
                </h2>

                <p>
                    Номер присваивается
                    автоматически только
                    при окончательном
                    сохранении.
                </p>

                {% if certificate_doc
                    and
                    certificate_doc.status
                    == 'issued' %}

                    <p>
                        <span
                            class="status green"
                        >
                            Выдана
                        </span>
                    </p>

                    <p>
                        Справка №
                        <strong>
                            {{
                                certificate_doc
                                .certificate_number
                            }}
                        </strong>
                        от
                        {{
                            certificate_doc
                            .issue_date
                        }}
                    </p>

                    <a
                        class="btn btn-primary"
                        href="{{ url_for(
                            'download_post_doc_v8',
                            document_id=
                                certificate_doc.id
                        ) }}"
                    >
                        Скачать DOCX
                    </a>

                {% elif
                    student.status
                    == 'enrolled' %}

                    {% if certificate_doc %}
                        <p>
                            <span
                                class="
                                    status yellow
                                "
                            >
                                Черновик
                            </span>
                        </p>
                    {% endif %}

                    <a
                        class="btn btn-primary"
                        href="{{ url_for(
                            'certificate_v8',
                            student_id=
                                student.id
                        ) }}"
                    >
                        {% if certificate_doc %}
                            Продолжить
                        {% else %}
                            Сформировать
                        {% endif %}
                    </a>

                {% endif %}
            </div>
        </div>

        <div class="card space">
            <h2 style="margin-top:0">
                Данные ученика
            </h2>

            <table>
                <tr>
                    <th>ФИО</th>
                    <td>
                        {{ student.last_name }}
                        {{ student.first_name }}
                        {{ student.middle_name or '' }}
                    </td>
                </tr>

                <tr>
                    <th>Дата рождения</th>
                    <td>
                        {{ student.birth_date }}
                    </td>
                </tr>

                <tr>
                    <th>Класс</th>
                    <td>
                        {{ student.class_number }}
                    </td>
                </tr>

                <tr>
                    <th>Учебный год</th>
                    <td>
                        {{ current_year }}
                    </td>
                </tr>
            </table>

            <a
                class="btn btn-secondary space"
                href="{{ url_for(
                    'edit_post_doc_student_v8',
                    student_id=student.id
                ) }}"
            >
                Редактировать данные
                ученика
            </a>
        </div>

        {% if current_user.role
            in ('attestation', 'admin') %}

            <div class="space">
                <a
                    href="{{ url_for(
                        'post_docs_registry_v8'
                    ) }}"
                >
                    Реестр сформированных
                    документов →
                </a>
            </div>

        {% endif %}

        {% if current_user.role
            == 'admin' %}

            <div class="space">
                <a
                    href="{{ url_for(
                        'post_docs_settings_v8'
                    ) }}"
                >
                    Настройки нумерации
                    и учебного года →
                </a>
            </div>

        {% endif %}
        """

        return render_page(
            "Документы после зачисления",
            body,
            student=student,
            order_doc=order_doc,
            certificate_doc=(
                certificate_doc
            ),
            current_year=current_year,
        )

    # ========================================================
    # Редактирование базовых данных ученика
    # даже после зачисления
    # ========================================================

    def edit_student_data(
        student_id: int,
    ):

        student = get_student_or_404(
            student_id
        )

        if request.method == "POST":

            last_name = request.form.get(
                "last_name",
                "",
            ).strip()

            first_name = request.form.get(
                "first_name",
                "",
            ).strip()

            middle_name = request.form.get(
                "middle_name",
                "",
            ).strip()

            birth_date_value = (
                request.form.get(
                    "birth_date",
                    "",
                ).strip()
            )

            class_value = request.form.get(
                "class_number",
                "",
            ).strip()

            if (
                not last_name
                or not first_name
                or not birth_date_value
                or not class_value.isdigit()
            ):
                flash(
                    "Заполните обязательные "
                    "поля.",
                    "error",
                )
                return redirect(
                    request.url
                )

            try:
                date.fromisoformat(
                    birth_date_value
                )
            except ValueError:
                flash(
                    "Некорректная дата "
                    "рождения.",
                    "error",
                )
                return redirect(
                    request.url
                )

            class_number = int(
                class_value
            )

            if not 1 <= class_number <= 11:
                flash(
                    "Класс должен быть "
                    "от 1 до 11.",
                    "error",
                )
                return redirect(
                    request.url
                )

            before = {
                "last_name":
                    student["last_name"],

                "first_name":
                    student["first_name"],

                "middle_name":
                    student["middle_name"],

                "birth_date":
                    student["birth_date"],

                "class_number":
                    student["class_number"],
            }

            get_db().execute(
                """
                UPDATE students
                SET
                    last_name = ?,
                    first_name = ?,
                    middle_name = ?,
                    birth_date = ?,
                    class_number = ?,
                    updated_at = ?
                WHERE id = ?
                """,
                (
                    last_name,
                    first_name,
                    middle_name,
                    birth_date_value,
                    class_number,
                    now_iso(),
                    student_id,
                ),
            )

            # Если справка еще только
            # черновик, удаляем ручную
            # форму ФИО, чтобы она заново
            # рассчиталась из обновленной
            # карточки.
            get_db().execute(
                """
                UPDATE
                    post_enrollment_documents
                SET
                    name_override = NULL,
                    updated_at = ?,
                    updated_by = ?
                WHERE student_id = ?
                  AND document_type =
                    'certificate'
                  AND status = 'draft'
                """,
                (
                    now_iso(),
                    g.current_user["id"],
                    student_id,
                ),
            )

            get_db().commit()

            after = {
                "last_name":
                    last_name,

                "first_name":
                    first_name,

                "middle_name":
                    middle_name,

                "birth_date":
                    birth_date_value,

                "class_number":
                    class_number,
            }

            audit(
                "post_doc_student_data_edited",
                student_id,
                json.dumps(
                    {
                        "before": before,
                        "after": after,
                    },
                    ensure_ascii=False,
                ),
            )

            flash(
                "Данные ученика "
                "обновлены.",
                "success",
            )

            return redirect(
                url_for(
                    "post_docs_v8",
                    student_id=student_id,
                )
            )

        body = """
        <h1>
            Редактировать данные ученика
        </h1>

        <div class="alert info">
            Изменения применятся к
            <strong>черновикам</strong>
            справки и выписки.

            Уже окончательно сформированные
            документы не изменяются.
        </div>

        <form
            class="form-section"
            method="post"
        >
            <input
                type="hidden"
                name="csrf_token"
                value="{{ csrf_token }}"
            >

            <div class="form-grid">

                <div>
                    <label>Фамилия *</label>
                    <input
                        name="last_name"
                        value="{{
                            student.last_name
                        }}"
                        required
                    >
                </div>

                <div>
                    <label>Имя *</label>
                    <input
                        name="first_name"
                        value="{{
                            student.first_name
                        }}"
                        required
                    >
                </div>

                <div>
                    <label>Отчество</label>
                    <input
                        name="middle_name"
                        value="{{
                            student.middle_name
                            or ''
                        }}"
                    >
                </div>

                <div>
                    <label>
                        Дата рождения *
                    </label>

                    <input
                        type="date"
                        name="birth_date"
                        value="{{
                            student.birth_date
                        }}"
                        required
                    >
                </div>

                <div>
                    <label>Класс *</label>

                    <select
                        name="class_number"
                        required
                    >
                        {% for number
                           in range(1, 12) %}
                            <option
                                value="{{ number }}"
                                {% if
                                    student.class_number
                                    == number
                                %}
                                    selected
                                {% endif %}
                            >
                                {{ number }}
                            </option>
                        {% endfor %}
                    </select>
                </div>
            </div>

            <div
                class="space"
                style="
                    display:flex;
                    gap:10px;
                    flex-wrap:wrap;
                "
            >
                <button
                    class="btn btn-primary"
                    type="submit"
                >
                    Сохранить изменения
                </button>

                <a
                    class="btn btn-secondary"
                    href="{{ url_for(
                        'post_docs_v8',
                        student_id=student.id
                    ) }}"
                >
                    Назад
                </a>
            </div>
        </form>
        """

        return render_page(
            "Редактирование ученика",
            body,
            student=student,
            range=range,
        )

    # ========================================================
    # Выписка: черновик и просмотр
    # ========================================================

    def order_extract_page(
        student_id: int,
    ):

        student = get_student_or_404(
            student_id
        )

        require_enrolled(
            student
        )

        document = get_post_document(
            student_id,
            "order_extract",
        )

        if (
            document
            and document["status"]
            == "issued"
        ):
            return redirect(
                url_for(
                    "post_docs_v8",
                    student_id=student_id,
                )
            )

        if request.method == "POST":

            order_number = request.form.get(
                "order_number",
                "",
            ).strip()

            order_date = request.form.get(
                "order_date",
                "",
            ).strip()

            if not order_number:
                flash(
                    "Введите номер приказа.",
                    "error",
                )
                return redirect(
                    request.url
                )

            try:
                date.fromisoformat(
                    order_date
                )
            except ValueError:
                flash(
                    "Введите корректную "
                    "дату приказа.",
                    "error",
                )
                return redirect(
                    request.url
                )

            timestamp = now_iso()

            if document:

                get_db().execute(
                    """
                    UPDATE
                        post_enrollment_documents
                    SET
                        order_number = ?,
                        order_date = ?,
                        updated_by = ?,
                        updated_at = ?
                    WHERE id = ?
                    """,
                    (
                        order_number,
                        order_date,
                        g.current_user["id"],
                        timestamp,
                        document["id"],
                    ),
                )

            else:

                get_db().execute(
                    """
                    INSERT INTO
                        post_enrollment_documents (
                            student_id,
                            document_type,
                            status,
                            order_number,
                            order_date,
                            created_by,
                            created_at,
                            updated_by,
                            updated_at
                        )
                    VALUES (
                        ?,
                        'order_extract',
                        'draft',
                        ?,
                        ?,
                        ?,
                        ?,
                        ?,
                        ?
                    )
                    """,
                    (
                        student_id,
                        order_number,
                        order_date,
                        g.current_user["id"],
                        timestamp,
                        g.current_user["id"],
                        timestamp,
                    ),
                )

            get_db().commit()

            audit(
                "order_extract_draft_saved",
                student_id,
                (
                    f"order_number="
                    f"{order_number}; "
                    f"order_date="
                    f"{order_date}"
                ),
            )

            return redirect(
                url_for(
                    "order_extract_v8",
                    student_id=student_id,
                    preview=1,
                )
            )

        document = get_post_document(
            student_id,
            "order_extract",
        )

        preview = (
            request.args.get(
                "preview"
            )
            == "1"
            and document is not None
        )

        body = """
        <h1>
            Выписка из приказа
        </h1>

        <div class="alert info">
            Введите реквизиты
            <strong>исходного приказа</strong>.

            Собственного номера
            выписка не получает.
        </div>

        <form
            class="form-section"
            method="post"
        >
            <input
                type="hidden"
                name="csrf_token"
                value="{{ csrf_token }}"
            >

            <div class="form-grid">

                <div>
                    <label>
                        Номер приказа *
                    </label>

                    <input
                        name="order_number"
                        value="{{
                            document.order_number
                            if document
                            else ''
                        }}"
                        placeholder="
                            1-Э-26/27
                        "
                        required
                    >
                </div>

                <div>
                    <label>
                        Дата приказа *
                    </label>

                    <input
                        type="date"
                        name="order_date"
                        value="{{
                            document.order_date
                            if document
                            else ''
                        }}"
                        required
                    >
                </div>
            </div>

            <button
                class="btn btn-primary space"
                type="submit"
            >
                Предварительный просмотр
            </button>
        </form>

        {% if preview %}

            <div class="card space">
                <h2 style="margin-top:0">
                    Предварительный просмотр
                </h2>

                <p>
                    <strong>
                        Выписка из приказа
                    </strong>
                </p>

                <p>
                    Приказ №
                    <strong>
                        {{ document.order_number }}
                    </strong>
                    от
                    <strong>
                        {{
                            short_date(
                                document.order_date
                            )
                        }}
                    </strong>
                </p>

                <p>
                    Обучающийся:
                    <strong>
                        {{ student_name }}
                    </strong>
                </p>

                <p>
                    Класс:
                    <strong>
                        {{ student.class_number }}
                    </strong>
                </p>

                <p>
                    Дата выписки:
                    <strong>
                        {{ today_text }}
                    </strong>
                </p>

                <div class="alert info">
                    Пока это черновик.
                    Изменения в карточке ученика
                    будут применены при повторном
                    открытии.
                </div>

                <div
                    style="
                        display:flex;
                        gap:10px;
                        flex-wrap:wrap;
                    "
                >
                    <a
                        class="btn btn-secondary"
                        href="{{ url_for(
                            'edit_post_doc_student_v8',
                            student_id=student.id
                        ) }}"
                    >
                        Исправить данные
                        ученика
                    </a>

                    <form
                        class="inline"
                        method="post"
                        action="{{ url_for(
                            'finalize_order_extract_v8',
                            student_id=student.id
                        ) }}"
                    >
                        <input
                            type="hidden"
                            name="csrf_token"
                            value="{{ csrf_token }}"
                        >

                        <button
                            class="btn btn-green"
                            type="submit"
                        >
                            Сохранить
                            окончательно
                        </button>
                    </form>
                </div>
            </div>

        {% endif %}

        <div class="space">
            <a
                class="btn btn-secondary"
                href="{{ url_for(
                    'post_docs_v8',
                    student_id=student.id
                ) }}"
            >
                Назад
            </a>
        </div>
        """

        return render_page(
            "Выписка из приказа",
            body,
            student=student,
            document=document,
            preview=preview,
            student_name=(
                student_full_name(
                    student
                )
            ),
            today_text=(
                short_date(
                    date.today()
                )
            ),
            short_date=short_date,
        )

    # ========================================================
    # Окончательное сохранение выписки
    # ========================================================

    def finalize_order_extract(
        student_id: int,
    ):

        student = get_student_or_404(
            student_id
        )

        require_enrolled(
            student
        )

        document = get_post_document(
            student_id,
            "order_extract",
        )

        if not document:
            abort(
                400,
                "Сначала заполните "
                "черновик выписки."
            )

        if (
            document["status"]
            == "issued"
        ):
            return redirect(
                url_for(
                    "post_docs_v8",
                    student_id=student_id,
                )
            )

        if (
            not document["order_number"]
            or not document["order_date"]
        ):
            abort(
                400,
                "Не заполнены реквизиты "
                "приказа."
            )

        issue_date = date.today()

        snapshot = {
            "student_name":
                student_full_name(
                    student
                ),

            "class_number":
                int(
                    student[
                        "class_number"
                    ]
                ),

            "order_number":
                document[
                    "order_number"
                ],

            "order_date":
                document[
                    "order_date"
                ],

            "issue_date":
                issue_date.isoformat(),

            "academic_year":
                academic_year(),
        }

        try:
            content = (
                generate_order_extract(
                    snapshot
                )
            )

        except Exception as error:
            flash(
                "Не удалось сформировать "
                f"выписку: {error}",
                "error",
            )

            return redirect(
                url_for(
                    "order_extract_v8",
                    student_id=student_id,
                    preview=1,
                )
            )

        stored_name = (
            "order_extract_"
            f"{student_id}_"
            f"{uuid.uuid4().hex}.docx"
        )

        path = (
            generated_dir
            / stored_name
        )

        path.write_bytes(
            content
        )

        timestamp = now_iso()

        get_db().execute(
            """
            UPDATE
                post_enrollment_documents
            SET
                status = 'issued',
                issue_date = ?,
                academic_year = ?,
                snapshot_json = ?,
                stored_name = ?,
                finalized_by = ?,
                finalized_at = ?,
                updated_by = ?,
                updated_at = ?
            WHERE id = ?
            """,
            (
                issue_date.isoformat(),
                academic_year(),
                json.dumps(
                    snapshot,
                    ensure_ascii=False,
                ),
                stored_name,
                g.current_user["id"],
                timestamp,
                g.current_user["id"],
                timestamp,
                document["id"],
            ),
        )

        get_db().commit()

        audit(
            "order_extract_issued",
            student_id,
            (
                f"order_number="
                f"{document['order_number']}; "
                f"order_date="
                f"{document['order_date']}"
            ),
        )

        flash(
            "Выписка сформирована "
            "и сохранена.",
            "success",
        )

        return redirect(
            url_for(
                "post_docs_v8",
                student_id=student_id,
            )
        )

    # ========================================================
    # Справка: черновик
    # ========================================================


    def certificate_page(
        student_id: int,
    ):

        student = get_student_or_404(
            student_id
        )

        require_enrolled(
            student
        )

        document = get_post_document(
            student_id,
            "certificate",
        )

        if (
            document
            and document["status"]
            == "issued"
        ):
            return redirect(
                url_for(
                    "post_docs_v8",
                    student_id=student_id,
                )
            )

        automatic_name = (
            automatic_certificate_name(
                student
            )
        )

        if request.method == "POST":

            name_form = request.form.get(
                "certificate_name",
                "",
            ).strip()

            if not name_form:
                name_form = automatic_name

            order_number = (
                request.form.get(
                    "order_number",
                    "",
                ).strip()
            )

            order_date_value = (
                request.form.get(
                    "order_date",
                    "",
                ).strip()
            )

            if not order_number:
                flash(
                    "Введите номер приказа "
                    "о зачислении.",
                    "error",
                )

                return redirect(
                    request.url
                )

            try:
                date.fromisoformat(
                    order_date_value
                )

            except ValueError:
                flash(
                    "Введите корректную "
                    "дату приказа.",
                    "error",
                )

                return redirect(
                    request.url
                )

            override = (
                None
                if name_form
                == automatic_name
                else name_form
            )

            timestamp = now_iso()

            if document:

                get_db().execute(
                    """
                    UPDATE
                        post_enrollment_documents
                    SET
                        name_override = ?,
                        order_number = ?,
                        order_date = ?,
                        updated_by = ?,
                        updated_at = ?
                    WHERE id = ?
                    """,
                    (
                        override,
                        order_number,
                        order_date_value,
                        g.current_user["id"],
                        timestamp,
                        document["id"],
                    ),
                )

            else:

                get_db().execute(
                    """
                    INSERT INTO
                        post_enrollment_documents (
                            student_id,
                            document_type,
                            status,
                            name_override,
                            order_number,
                            order_date,
                            created_by,
                            created_at,
                            updated_by,
                            updated_at
                        )
                    VALUES (
                        ?,
                        'certificate',
                        'draft',
                        ?,
                        ?,
                        ?,
                        ?,
                        ?,
                        ?,
                        ?
                    )
                    """,
                    (
                        student_id,
                        override,
                        order_number,
                        order_date_value,
                        g.current_user["id"],
                        timestamp,
                        g.current_user["id"],
                        timestamp,
                    ),
                )

            get_db().commit()

            audit(
                "certificate_draft_saved",
                student_id,
                (
                    f"order_number="
                    f"{order_number}; "
                    f"order_date="
                    f"{order_date_value}"
                ),
            )

            return redirect(
                url_for(
                    "certificate_v8",
                    student_id=student_id,
                    preview=1,
                )
            )

        document = get_post_document(
            student_id,
            "certificate",
        )

        certificate_name = (
            document["name_override"]
            if (
                document
                and
                document["name_override"]
            )
            else automatic_name
        )

        preview = (
            request.args.get(
                "preview"
            )
            == "1"
            and document is not None
        )

        body = """
        <h1>
            Справка
        </h1>

        <div class="alert info">
            Номер справки будет присвоен
            автоматически только после
            окончательного сохранения.

            <br><br>

            Номер и дата приказа
            о зачислении вводятся
            вручную и сохраняются
            в черновике справки.
        </div>

        <form
            class="form-section"
            method="post"
        >
            <input
                type="hidden"
                name="csrf_token"
                value="{{ csrf_token }}"
            >

            <div class="form-grid">

                <div>
                    <label>
                        ФИО для текста
                        справки *
                    </label>

                    <input
                        name="certificate_name"
                        value="{{
                            certificate_name
                        }}"
                        required
                    >

                    <p class="muted">
                        Форма для фразы
                        «Дана кому».
                        При необходимости
                        исправьте вручную.
                    </p>
                </div>

                <div>
                    <label>
                        Номер приказа
                        о зачислении *
                    </label>

                    <input
                        name="order_number"
                        value="{{
                            document.order_number
                            if document
                            and
                            document.order_number
                            else ''
                        }}"
                        placeholder="
                            1-Э-26/27
                        "
                        required
                    >
                </div>

                <div>
                    <label>
                        Дата приказа
                        о зачислении *
                    </label>

                    <input
                        type="date"
                        name="order_date"
                        value="{{
                            document.order_date
                            if document
                            and
                            document.order_date
                            else ''
                        }}"
                        required
                    >
                </div>
            </div>

            <button
                class="btn btn-primary space"
                type="submit"
            >
                Предварительный просмотр
            </button>
        </form>

        {% if preview %}

            <div class="card space">

                <h2 style="margin-top:0">
                    Предварительный
                    просмотр
                </h2>

                <p>
                    <strong>
                        Справка № —
                    </strong>

                    номер будет присвоен
                    при окончательном
                    сохранении
                </p>

                <table>
                    <tr>
                        <th>
                            ФИО
                        </th>

                        <td>
                            {{
                                certificate_name
                            }}
                        </td>
                    </tr>

                    <tr>
                        <th>
                            Дата рождения
                        </th>

                        <td>
                            {{
                                short_date(
                                    student.birth_date
                                )
                            }}
                        </td>
                    </tr>

                    <tr>
                        <th>
                            Класс
                        </th>

                        <td>
                            {{
                                student.class_number
                            }}
                        </td>
                    </tr>

                    <tr>
                        <th>
                            Учебный год
                        </th>

                        <td>
                            {{
                                current_year
                            }}
                        </td>
                    </tr>

                    <tr>
                        <th>
                            Приказ
                        </th>

                        <td>
                            №
                            {{
                                document.order_number
                            }}
                            от
                            {{
                                short_date(
                                    document.order_date
                                )
                            }}
                        </td>
                    </tr>

                    <tr>
                        <th>
                            Дата справки
                        </th>

                        <td>
                            {{
                                today_text
                            }}
                        </td>
                    </tr>
                </table>

                <div class="alert info space">
                    Это черновик.

                    Номер справки еще
                    не занят.

                    <br><br>

                    Если исправить данные
                    ученика и открыть
                    справку снова,
                    ФИО, дата рождения
                    и класс подтянутся
                    из карточки заново.

                    <br><br>

                    Номер и дата приказа
                    останутся сохраненными
                    в черновике.
                </div>

                <div
                    style="
                        display:flex;
                        gap:10px;
                        flex-wrap:wrap;
                    "
                >

                    <a
                        class="btn btn-secondary"
                        href="{{ url_for(
                            'edit_post_doc_student_v8',
                            student_id=student.id
                        ) }}"
                    >
                        Исправить данные
                        ученика
                    </a>

                    <form
                        class="inline"
                        method="post"
                        action="{{ url_for(
                            'finalize_certificate_v8',
                            student_id=student.id
                        ) }}"
                    >
                        <input
                            type="hidden"
                            name="csrf_token"
                            value="{{ csrf_token }}"
                        >

                        <button
                            class="btn btn-green"
                            type="submit"
                        >
                            Сохранить окончательно
                            и присвоить номер
                        </button>
                    </form>
                </div>
            </div>

        {% endif %}

        <div class="space">

            <a
                class="btn btn-secondary"
                href="{{ url_for(
                    'post_docs_v8',
                    student_id=student.id
                ) }}"
            >
                Назад
            </a>

        </div>
        """

        return render_page(
            "Справка",
            body,
            student=student,
            document=document,
            certificate_name=(
                certificate_name
            ),
            preview=preview,
            short_date=short_date,
            current_year=academic_year(),
            today_text=short_date(
                date.today()
            ),
        )


    # ========================================================
    # Окончательная регистрация справки
    # ========================================================


    def finalize_certificate(
        student_id: int,
    ):

        student = get_student_or_404(
            student_id
        )

        require_enrolled(
            student
        )

        db = get_db()

        generated_path = None

        try:

            # Сквозной номер получаем
            # только на этом этапе.
            db.execute(
                "BEGIN IMMEDIATE"
            )

            document = db.execute(
                """
                SELECT *
                FROM post_enrollment_documents
                WHERE student_id = ?
                  AND document_type =
                    'certificate'
                """,
                (student_id,),
            ).fetchone()

            if not document:
                raise RuntimeError(
                    "Сначала откройте "
                    "предварительный просмотр "
                    "справки."
                )

            if (
                document["status"]
                == "issued"
            ):
                db.rollback()

                return redirect(
                    url_for(
                        "post_docs_v8",
                        student_id=student_id,
                    )
                )

            if not document[
                "order_number"
            ]:
                raise RuntimeError(
                    "Не указан номер "
                    "приказа о зачислении."
                )

            if not document[
                "order_date"
            ]:
                raise RuntimeError(
                    "Не указана дата "
                    "приказа о зачислении."
                )

            counter = db.execute(
                """
                SELECT last_value
                FROM document_counters
                WHERE counter_key =
                    'certificate'
                """
            ).fetchone()

            last_value = (
                int(counter[0])
                if counter
                else 0
            )

            next_number = (
                last_value + 1
            )

            current_year = (
                academic_year()
            )

            # Дата справки —
            # дата нажатия
            # «Сохранить окончательно».
            issue_date = date.today()

            certificate_name = (
                document[
                    "name_override"
                ]
                or
                automatic_certificate_name(
                    student
                )
            )

            snapshot = {

                "certificate_number":
                    next_number,

                "certificate_name":
                    certificate_name,

                "student_name":
                    student_full_name(
                        student
                    ),

                "birth_date":
                    student[
                        "birth_date"
                    ],

                "class_number":
                    int(
                        student[
                            "class_number"
                        ]
                    ),

                "academic_year":
                    current_year,

                "order_number":
                    document[
                        "order_number"
                    ],

                "order_date":
                    document[
                        "order_date"
                    ],

                "issue_date":
                    issue_date.isoformat(),
            }

            content = (
                generate_certificate(
                    snapshot
                )
            )

            stored_name = (
                "certificate_"
                f"{next_number}_"
                f"{student_id}_"
                f"{uuid.uuid4().hex}"
                ".docx"
            )

            generated_path = (
                generated_dir
                / stored_name
            )

            generated_path.write_bytes(
                content
            )

            timestamp = now_iso()

            db.execute(
                """
                UPDATE document_counters
                SET last_value = ?
                WHERE counter_key =
                    'certificate'
                """,
                (
                    next_number,
                ),
            )

            db.execute(
                """
                UPDATE
                    post_enrollment_documents
                SET
                    status = 'issued',
                    certificate_number = ?,
                    issue_date = ?,
                    academic_year = ?,
                    snapshot_json = ?,
                    stored_name = ?,
                    finalized_by = ?,
                    finalized_at = ?,
                    updated_by = ?,
                    updated_at = ?
                WHERE id = ?
                """,
                (
                    next_number,
                    issue_date.isoformat(),
                    current_year,

                    json.dumps(
                        snapshot,
                        ensure_ascii=False,
                    ),

                    stored_name,

                    g.current_user["id"],
                    timestamp,

                    g.current_user["id"],
                    timestamp,

                    document["id"],
                ),
            )

            db.commit()

        except Exception as error:

            try:
                db.rollback()

            except Exception:
                pass

            if (
                generated_path
                and
                generated_path.exists()
            ):
                try:
                    generated_path.unlink()

                except OSError:
                    pass

            flash(
                "Не удалось сформировать "
                f"справку: {error}",
                "error",
            )

            return redirect(
                url_for(
                    "certificate_v8",
                    student_id=student_id,
                    preview=1,
                )
            )

        audit(
            "certificate_issued",
            student_id,
            (
                f"certificate_number="
                f"{next_number}; "
                f"order_number="
                f"{document['order_number']}; "
                f"order_date="
                f"{document['order_date']}"
            ),
        )

        flash(
            f"Справка № {next_number} "
            "сформирована и "
            "зарегистрирована.",
            "success",
        )

        return redirect(
            url_for(
                "post_docs_v8",
                student_id=student_id,
            )
        )


    # ========================================================
    # Скачивание
    # ========================================================

    def download_post_document(
        document_id: int,
    ):

        document = get_db().execute(
            """
            SELECT *
            FROM post_enrollment_documents
            WHERE id = ?
            """,
            (document_id,),
        ).fetchone()

        if not document:
            abort(404)

        student = get_student_or_404(
            document["student_id"]
        )

        if (
            document["status"]
            != "issued"
            or not document["stored_name"]
        ):
            abort(
                400,
                "Черновик нельзя скачать "
                "как окончательный документ."
            )

        path = (
            generated_dir
            / document["stored_name"]
        )

        if not path.exists():
            abort(
                404,
                "Файл документа "
                "не найден."
            )

        if (
            document["document_type"]
            == "certificate"
        ):
            download_name = (
                "Справка № "
                f"{document['certificate_number']} "
                f"{student_full_name(student)}"
                ".docx"
            )

        else:
            download_name = (
                "Выписка из приказа "
                f"{student_full_name(student)}"
                ".docx"
            )

        audit(
            "post_doc_downloaded",
            student["id"],
            (
                f"document_id="
                f"{document_id}; "
                f"type="
                f"{document['document_type']}"
            ),
        )

        return send_file(
            path,
            as_attachment=True,
            download_name=download_name,
            mimetype=(
                "application/"
                "vnd.openxmlformats-"
                "officedocument."
                "wordprocessingml.document"
            ),
        )

    # ========================================================
    # Реестр
    # ========================================================

    def registry_page():

        query = """
            SELECT
                d.*,
                s.last_name,
                s.first_name,
                s.middle_name,
                s.class_number,
                b.name AS branch_name
            FROM post_enrollment_documents d
            JOIN students s
              ON s.id = d.student_id
            JOIN branches b
              ON b.id = s.branch_id
            WHERE d.status = 'issued'
        """

        params = []

        document_type = (
            request.args.get(
                "document_type",
                "",
            ).strip()
        )

        search = request.args.get(
            "search",
            "",
        ).strip()

        if document_type in {
            "order_extract",
            "certificate",
        }:
            query += (
                " AND d.document_type = ?"
            )
            params.append(
                document_type
            )

        if search:

            pattern = (
                f"%{search}%"
            )

            query += """
                AND (
                    s.last_name LIKE ?
                    OR s.first_name LIKE ?
                    OR CAST(
                        d.certificate_number
                        AS TEXT
                    ) LIKE ?
                    OR d.order_number LIKE ?
                )
            """

            params.extend(
                [
                    pattern,
                    pattern,
                    pattern,
                    pattern,
                ]
            )

        query += (
            " ORDER BY "
            "d.finalized_at DESC"
        )

        rows = get_db().execute(
            query,
            params,
        ).fetchall()

        body = """
        <h1>
            Реестр сформированных
            документов
        </h1>

        <form
            class="card form-grid"
            method="get"
        >

            <div>
                <label>Поиск</label>

                <input
                    name="search"
                    value="{{
                        request.args.get(
                            'search',
                            ''
                        )
                    }}"
                    placeholder="
                        ФИО, номер справки
                        или приказа
                    "
                >
            </div>

            <div>
                <label>
                    Тип документа
                </label>

                <select
                    name="document_type"
                >
                    <option value="">
                        Все
                    </option>

                    <option
                        value="certificate"
                    >
                        Справка
                    </option>

                    <option
                        value="order_extract"
                    >
                        Выписка из приказа
                    </option>
                </select>
            </div>

            <div style="align-self:end">
                <button
                    class="btn btn-primary"
                >
                    Найти
                </button>
            </div>
        </form>

        <div class="card space">

            <table>
                <thead>
                    <tr>
                        <th>
                            Документ
                        </th>
                        <th>Ученик</th>
                        <th>Филиал</th>
                        <th>Класс</th>
                        <th>
                            Дата
                        </th>
                        <th></th>
                    </tr>
                </thead>

                <tbody>
                    {% for row in rows %}
                        <tr>
                            <td>
                                {% if
                                    row.document_type
                                    == 'certificate'
                                %}
                                    <strong>
                                        Справка №
                                        {{
                                            row
                                            .certificate_number
                                        }}
                                    </strong>
                                {% else %}
                                    <strong>
                                        Выписка
                                    </strong>
                                    <br>
                                    Приказ №
                                    {{
                                        row
                                        .order_number
                                    }}
                                    от
                                    {{
                                        row
                                        .order_date
                                    }}
                                {% endif %}
                            </td>

                            <td>
                                <a
                                    href="{{ url_for(
                                        'student_detail',
                                        student_id=
                                            row.student_id
                                    ) }}"
                                >
                                    {{
                                        row.last_name
                                    }}
                                    {{
                                        row.first_name
                                    }}
                                    {{
                                        row.middle_name
                                        or ''
                                    }}
                                </a>
                            </td>

                            <td>
                                {{
                                    row.branch_name
                                }}
                            </td>

                            <td>
                                {{
                                    row.class_number
                                }}
                            </td>

                            <td>
                                {{
                                    row.issue_date
                                }}
                            </td>

                            <td>
                                <a
                                    href="{{ url_for(
                                        'download_post_doc_v8',
                                        document_id=
                                            row.id
                                    ) }}"
                                >
                                    Скачать
                                </a>
                            </td>
                        </tr>

                    {% else %}

                        <tr>
                            <td colspan="6">
                                Документов
                                пока нет.
                            </td>
                        </tr>

                    {% endfor %}
                </tbody>
            </table>
        </div>
        """

        return render_page(
            "Реестр документов",
            body,
            rows=rows,
        )

    # ========================================================
    # Настройки
    # ========================================================

    def settings_page():

        db = get_db()

        maximum = db.execute(
            """
            SELECT COALESCE(
                MAX(certificate_number),
                0
            )
            FROM post_enrollment_documents
            WHERE document_type =
                'certificate'
            """
        ).fetchone()[0]

        counter = db.execute(
            """
            SELECT last_value
            FROM document_counters
            WHERE counter_key =
                'certificate'
            """
        ).fetchone()[0]

        if request.method == "POST":

            year_value = request.form.get(
                "academic_year",
                "",
            ).strip()

            counter_value = (
                request.form.get(
                    "last_certificate_number",
                    "",
                ).strip()
            )

            if not re.fullmatch(
                r"\d{4}-\d{4}",
                year_value,
            ):
                flash(
                    "Учебный год должен "
                    "быть в формате "
                    "2026-2027.",
                    "error",
                )
                return redirect(
                    request.url
                )

            if not counter_value.isdigit():
                flash(
                    "Последний номер "
                    "справки должен быть "
                    "целым числом.",
                    "error",
                )
                return redirect(
                    request.url
                )

            counter_number = int(
                counter_value
            )

            if counter_number < int(
                maximum
            ):
                flash(
                    "Нельзя установить "
                    "счетчик ниже уже "
                    "выданного номера "
                    f"{maximum}.",
                    "error",
                )
                return redirect(
                    request.url
                )

            db.execute(
                """
                INSERT INTO portal_settings (
                    setting_key,
                    setting_value
                )
                VALUES (
                    'academic_year',
                    ?
                )
                ON CONFLICT(setting_key)
                DO UPDATE SET
                    setting_value =
                    excluded.setting_value
                """,
                (year_value,),
            )

            db.execute(
                """
                UPDATE document_counters
                SET last_value = ?
                WHERE counter_key =
                    'certificate'
                """,
                (counter_number,),
            )

            db.commit()

            audit(
                "post_docs_settings_updated",
                details=(
                    f"academic_year="
                    f"{year_value}; "
                    f"certificate_counter="
                    f"{counter_number}"
                ),
            )

            flash(
                "Настройки сохранены.",
                "success",
            )

            return redirect(
                request.url
            )

        body = """
        <h1>
            Настройки документов
            после зачисления
        </h1>

        <div class="alert info">
            Если до запуска портала
            справки уже нумеровались
            вручную, укажите здесь
            <strong>
                последний использованный
                номер
            </strong>.

            Следующая справка получит
            номер +1.
        </div>

        <form
            class="form-section"
            method="post"
        >
            <input
                type="hidden"
                name="csrf_token"
                value="{{ csrf_token }}"
            >

            <div class="form-grid">

                <div>
                    <label>
                        Учебный год
                    </label>

                    <input
                        name="academic_year"
                        value="{{
                            current_year
                        }}"
                        placeholder="
                            2026-2027
                        "
                        required
                    >
                </div>

                <div>
                    <label>
                        Последний
                        использованный
                        номер справки
                    </label>

                    <input
                        type="number"
                        min="{{ maximum }}"
                        name="
                            last_certificate_number
                        "
                        value="{{ counter }}"
                        required
                    >
                </div>
            </div>

            <button
                class="btn btn-primary space"
                type="submit"
            >
                Сохранить настройки
            </button>
        </form>
        """

        return render_page(
            "Настройки документов",
            body,
            current_year=academic_year(),
            counter=counter,
            maximum=maximum,
        )

    # ========================================================
    # Регистрация URL
    # ========================================================

    if (
        "post_docs_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/students/"
            "<int:student_id>/post-docs",
            endpoint="post_docs_v8",
            view_func=login_required(
                post_docs_page
            ),
            methods=["GET"],
        )

    if (
        "edit_post_doc_student_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/students/"
            "<int:student_id>/post-docs/edit",
            endpoint=
                "edit_post_doc_student_v8",
            view_func=login_required(
                edit_student_data
            ),
            methods=[
                "GET",
                "POST",
            ],
        )

    if (
        "order_extract_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/students/"
            "<int:student_id>/post-docs/"
            "order-extract",
            endpoint=
                "order_extract_v8",
            view_func=login_required(
                order_extract_page
            ),
            methods=[
                "GET",
                "POST",
            ],
        )

    if (
        "finalize_order_extract_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/students/"
            "<int:student_id>/post-docs/"
            "order-extract/finalize",
            endpoint=
                "finalize_order_extract_v8",
            view_func=login_required(
                finalize_order_extract
            ),
            methods=["POST"],
        )

    if (
        "certificate_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/students/"
            "<int:student_id>/post-docs/"
            "certificate",
            endpoint=
                "certificate_v8",
            view_func=login_required(
                certificate_page
            ),
            methods=[
                "GET",
                "POST",
            ],
        )

    if (
        "finalize_certificate_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/students/"
            "<int:student_id>/post-docs/"
            "certificate/finalize",
            endpoint=
                "finalize_certificate_v8",
            view_func=login_required(
                finalize_certificate
            ),
            methods=["POST"],
        )

    if (
        "download_post_doc_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/post-documents/"
            "<int:document_id>/download",
            endpoint=
                "download_post_doc_v8",
            view_func=login_required(
                download_post_document
            ),
            methods=["GET"],
        )

    if (
        "post_docs_registry_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/post-documents/registry",
            endpoint=
                "post_docs_registry_v8",
            view_func=roles_required(
                "attestation",
                "admin",
            )(
                registry_page
            ),
            methods=["GET"],
        )

    if (
        "post_docs_settings_v8"
        not in app.view_functions
    ):
        app.add_url_rule(
            "/post-documents/settings",
            endpoint=
                "post_docs_settings_v8",
            view_func=roles_required(
                "admin"
            )(
                settings_page
            ),
            methods=[
                "GET",
                "POST",
            ],
        )
